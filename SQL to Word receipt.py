import mysql.connector
from mysql.connector import Error
from docx import Document
import logging

def get_database_connection():
    try:
        connection = mysql.connector.connect(
            host='your_host',
            user='your_username',
            password='your_password',
            database='client_database'
        )
        if connection.is_connected():
            return connection
    except Error as e:
        logging.error(f"Error connecting to the database: {e}")
    return None

def get_clients_data(connection):
    clients_data = []
    if connection is not None:
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT name, address, oib_croatia FROM clients")
                clients_data = cursor.fetchall()
        except Error as e:
            logging.error(f"Error executing the query: {e}")
    return clients_data

def create_word_doc(clients_data):
    doc = Document()
    doc.add_heading('Clients Data', level=1)

    for idx, client in enumerate(clients_data, start=1):
        doc.add_heading(f'Client {idx}', level=2)
        doc.add_paragraph(f'Name: {client[0]}')
        doc.add_paragraph(f'Address: {client[1]}')
        doc.add_paragraph(f'OIB Croatia: {client[2]}')
        doc.add_page_break()

    try:
        doc.save('clients_data.docx')
        print('Word document "clients_data.docx" created successfully.')
    except Exception as e:
        logging.error(f"Error saving the Word document: {e}")

def main():
    logging.basicConfig(level=logging.INFO)  # Set the log level as needed

    connection = get_database_connection()
    if connection is not None:
        clients_data = get_clients_data(connection)
        create_word_doc(clients_data)
        connection.close()

if __name__ == '__main__':
    main()